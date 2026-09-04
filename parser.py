"""Normalize varied DOCX question layouts into the dashboard's existing JSON schema."""
import html
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches

# Accept a normal question header as well as labels such as
# "Question 3 (Non-Clinical)".  The parenthetical label is metadata, not part
# of the question text; the following paragraph remains the question text.
QUESTION_RE = re.compile(r"^(?:question|ques|q)\s*(\d+)(?:\s*\([^)]*\))?(?:\s*[:.)\-]\s*(.*)|\s*)$|^(\d+)\s*[.)]\s+(.+)$", re.I)
OPTION_RE = re.compile(r"^(?:option\s*)?([A-Ha-h]|[1-8])\s*[.)\]:\-]\s*(.+)$", re.I)
ANSWER_RE = re.compile(r"^(?:(?:correct\s*)?(?:answer|ans)|correct\s*(?:option|choice)?|right\s*answer)\s*(?:is)?\s*[:\-]?\s*(?:option|choice)?\s*\(?\s*([A-Ha-h]|[1-8])\s*\)?(?:[.)\]:\-]|\b)", re.I)
EXPLANATION_RE = re.compile(r"^(?:(?:brief|detailed)?\s*explanations?(?:\s+of\s+all\s+options)?|rationale|reason)\s*[:\-]?\s*(.*)$", re.I)
BRIEF_EXPLANATION_RE = re.compile(r"^brief\s*explanation\s*[:\-]?\s*(.*)$", re.I)
DETAILED_EXPLANATION_RE = re.compile(r"^detailed\s*explanation(?:\s+of\s+all\s+options)?\s*[:\-]?\s*(.*)$", re.I)
IMAGE_DESCRIPTION_RE = re.compile(r"^(?:image|figure|diagram|table)\s*description(?:\s+for\b[^:]*)?\s*[:\-]?\s*(.*)$", re.I)
# Plural "Explanations:" conventionally introduces the option-by-option section.
# Singular "Explanation:" is the normal explanation for the question.
EXPLANATIONS_HEADER_RE = re.compile(r"^(?:detailed\s+)?explanations(?:\s+of\s+all\s+options)?\s*[:\-]?\s*(.*)$", re.I)
OPTIONS_HEADER_RE = re.compile(r"^(?:options?|choices?)\s*[:\-]?\s*(.*)$", re.I)


def _clean(value):
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def _plain_docx_markup(value):
    """Remove Word-exported Markdown wrappers before classifying a line."""
    value = _clean(value)
    value = re.sub(r"^#{1,6}\s*", "", value)
    value = re.sub(r"^(?:[-•]\s+)(?=(?:\*\*)?Option\s+[A-H])", "", value, flags=re.I)
    # Some source documents use **Correct Answer:** and **Options:** as
    # literal text rather than Word bold.  They are still ordinary headings.
    value = value.replace("**", "").replace("*", "")
    if re.match(r"^\[\s*(?:image|figure|diagram|table)\s*description\b", value, re.I) and value.endswith("]"):
        value = value[1:-1].strip()
    return value


def _paragraph_items(document):
    """Read paragraphs and table cells in their original document order."""
    items = []
    def add_paragraph(paragraph, in_table=False):
        # A Word paragraph can contain hard line breaks.  Those breaks are
        # meaningful in question files: a common pattern is "Correct Answer:
        # B" followed by "Explanation: ..." in the *same* paragraph.
        lines = [_clean(line) for line in paragraph.text.splitlines() if _clean(line)]
        image_parts = []
        for blip in paragraph._p.xpath(".//a:blip"):
            relation_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if relation_id and relation_id in paragraph.part.related_parts:
                image_parts.append(paragraph.part.related_parts[relation_id])
        if lines or image_parts:
            style = paragraph.style.name.lower() if paragraph.style else ""
            for line_index, text in enumerate(lines or [""]):
                items.append({"text": text, "image_parts": image_parts if line_index == 0 else [], "in_table": in_table, "is_list": "list" in style or "bullet" in style, "is_heading": "heading" in style})
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            add_paragraph(Paragraph(child, document))
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        add_paragraph(paragraph, in_table=True)
    return items


def _save_image(part, image_dir, sequence):
    if not part or not image_dir:
        return None
    image_dir = Path(image_dir); image_dir.mkdir(parents=True, exist_ok=True)
    extension = {"image/png": ".png", "image/jpeg": ".jpg", "image/gif": ".gif", "image/bmp": ".bmp"}.get(part.content_type, ".png")
    path = image_dir / f"question-image-{sequence}{extension}"
    path.write_bytes(part.blob)
    return str(path)


def _candidate(number, text):
    return {"number": number, "text": text, "options": [], "answer_key": None, "explanation": [], "brief_explanation": [], "detailed_explanation": [], "detailed_heading": "Detailed Explanation of All Options:", "image_paths": [], "image_description": []}


def _paragraph_html(text):
    return f"<p>{html.escape(text)}</p>"


def _explanation_html(candidate):
    parts = []
    if candidate["explanation"]:
        parts.append("<p><strong>Explanation:</strong></p>")
        parts.extend(_paragraph_html(text) for text in candidate["explanation"])
    if candidate["brief_explanation"]:
        parts.append("<p><strong>Brief Explanation:</strong></p>")
        parts.extend(_paragraph_html(text) for text in candidate["brief_explanation"])
    if candidate["detailed_explanation"]:
        parts.append(f"<p><strong>{html.escape(candidate['detailed_heading'])}</strong></p><ul>")
        for text in candidate["detailed_explanation"]:
            match = re.match(r"^(Option\s+[A-H](?:\s*\([^)]*\))?(?:\s+is\s+(?:correct|incorrect))?\s*:)\s*(.*)$", text, re.I)
            if match:
                parts.append(f"<li><strong>{html.escape(match.group(1))}</strong> {html.escape(match.group(2))}</li>")
            else:
                parts.append(f"<li>{html.escape(text)}</li>")
        parts.append("</ul>")
    return "\n".join(parts)


def _add_inline_options(candidate, text):
    """Extract A. text B. text style option blocks written in one paragraph."""
    matches = re.findall(r"(?:^|\s)([A-H])\.\s*(.*?)(?=\s+[A-H]\.|$)", text, re.I)
    for raw_label, option_text in matches:
        label = raw_label.upper()
        if not any(existing == label for existing, _, _ in candidate["options"]):
            candidate["options"].append((label, len(candidate["options"]) + 1, _clean(option_text)))
    return bool(matches)


def _is_supported_numbered_question(items, index):
    """Avoid mistaking section headings and reference lists for numbered questions."""
    for next_item in items[index + 1:index + 9]:
        next_text = next_item["text"]
        if QUESTION_RE.match(next_text):
            return False
        if OPTION_RE.match(next_text):
            return True
    return False


def _question_json(candidate, subject_id, created_by):
    answer_value = next(value for label, value, _ in candidate["options"] if label == candidate["answer_key"])
    question_html = _paragraph_html(candidate["text"])
    if candidate["image_description"]:
        question_html += "<p><strong>Image Description:</strong> " + html.escape(" ".join(candidate["image_description"])) + "</p>"
    return {"uuid": str(uuid.uuid4()), "Question": question_html, "type": "SINGLE", "options": [{"name": f"<p>{html.escape(text)}</p>", "value": value, "imgUrl": "", "upload": False} for _, value, text in candidate["options"]], "Answer": {"options": answer_value}, "Question_image": "", "description": _explanation_html(candidate), "difficulty": "EASY", "previous_apperance": "", "perals": [], "tags": None, "questionId": "", "flags": {"pro": False, "editable": True, "qBank": True, "active": True, "testSeries": True, "paid": True}, "testSeries": None, "Question_order": candidate["number"], "syllabus": [subject_id] if subject_id else [], "math_library": "no", "description_image": "", "createdOn": datetime.now(timezone.utc).isoformat(), "createdBy": created_by, "modifiedOn": None}


def parse_docx(file_or_path, subject_id="", created_by=None, image_dir=None):
    """Detect common question styles; return only validated questions plus clear errors."""
    items = _paragraph_items(Document(file_or_path))
    candidates, current, mode, image_index = [], None, None, 0
    for index, item in enumerate(items):
        text = _plain_docx_markup(item["text"])
        question = QUESTION_RE.match(text) if text else None
        if question:
            # A bare "1." may be a section heading or bibliography item. Only accept it
            # when option lines follow; explicit Question/Q labels are always questions.
            explicit_question = bool(re.match(r"^(?:question|ques|q)\s*\d+\b", text, re.I))
            if not explicit_question and not _is_supported_numbered_question(items, index):
                continue
            if current: candidates.append(current)
            number = question.group(1) or question.group(3) or len(candidates) + 1
            current = _candidate(int(number), _clean(question.group(2) or question.group(4) or "")); mode = "question"
            continue
        if not current:
            continue
        if item["image_parts"]:
            for part in item["image_parts"]:
                image_index += 1; current["image_paths"].append(_save_image(part, image_dir, image_index))
            mode = "image"
            if text: current["image_description"].append(text)
            continue
        image_description = IMAGE_DESCRIPTION_RE.match(text) if text else None
        if image_description:
            mode = "image"; description = _clean(image_description.group(1))
            if description: current["image_description"].append(description)
            continue
        explanations_header = EXPLANATIONS_HEADER_RE.match(text) if text else None
        if explanations_header:
            mode = "detailed"
            current["detailed_heading"] = text.split(":", 1)[0].strip() + ":"
            first = _clean(explanations_header.group(1))
            if first: current["detailed_explanation"].append(first)
            continue
        detailed = DETAILED_EXPLANATION_RE.match(text) if text else None
        if detailed:
            mode = "detailed"; first = _clean(detailed.group(1))
            if first: current["detailed_explanation"].append(first)
            continue
        brief = BRIEF_EXPLANATION_RE.match(text) if text else None
        if brief:
            mode = "brief"; first = _clean(brief.group(1))
            if first: current["brief_explanation"].append(first)
            continue
        explanation = EXPLANATION_RE.match(text) if text else None
        if explanation:
            # A plain "Explanation:" belongs to the question itself.  Keep it
            # distinct from explicitly labelled Brief/Detailed sections.
            mode = "post_answer"; first = _clean(explanation.group(1))
            if first: current["explanation"].append(first)
            continue
        # Once the correct answer is read, no later text can be another answer
        # option.  It is explanation content until the next question header.
        if mode == "post_answer":
            current["explanation"].append(text)
            continue
        if mode == "brief":
            current["brief_explanation"].append(text)
            continue
        if mode == "detailed":
            current["detailed_explanation"].append(text)
            continue
        # Avoid treating "Option A is incorrect:" in a detailed explanation as
        # an Options: heading or as an additional multiple-choice answer.
        if mode == "detailed" and re.match(r"^Option\s+[A-H]\b", text, re.I):
            current["detailed_explanation"].append(text)
            continue
        options_header = OPTIONS_HEADER_RE.match(text) if text else None
        if options_header:
            mode = "options"; _add_inline_options(current, options_header.group(1)); continue
        option = OPTION_RE.match(text) if text else None
        if option:
            raw_label, option_text = option.groups()
            label = chr(64 + int(raw_label)) if raw_label.isdigit() else raw_label.upper()
            current["options"].append((label, len(current["options"]) + 1, _clean(option_text))); mode = "options"; continue
        answer = ANSWER_RE.match(text) if text else None
        if answer:
            raw_label = answer.group(1); current["answer_key"] = chr(64 + int(raw_label)) if raw_label.isdigit() else raw_label.upper(); mode = "post_answer"; continue
        if text:
            if mode == "options" and _add_inline_options(current, text):
                continue
            if mode == "post_answer": current["explanation"].append(text)
            elif mode == "brief": current["brief_explanation"].append(text)
            elif mode == "detailed": current["detailed_explanation"].append(text)
            elif mode == "image": current["image_description"].append(text)
            elif mode == "question": current["text"] = _clean(f"{current['text']} {text}")
    if current: candidates.append(current)

    created_by = created_by or {"uuid": "290604b4-498f-4428-8e4c-92cb9f3240d0", "name": "Akhil"}
    questions, errors, warnings = [], [], []
    for candidate in candidates:
        issues = []
        if not candidate["text"]: issues.append(("Question text", "Add text after the question number."))
        labels = [label for label, _, _ in candidate["options"]]
        expected_labels = [chr(65 + index) for index in range(len(labels))]
        if len(candidate["options"]) < 2 or labels != expected_labels: issues.append(("Options", "Provide two to eight sequentially labelled options (A, B, C and so on)."))
        if not candidate["answer_key"] or candidate["answer_key"] not in labels: issues.append(("Correct answer", "Add Answer: followed by one of the option letters."))
        if issues:
            errors.append({"question_number": candidate["number"], "question": candidate["text"] or "Unidentified question", "reason": "; ".join(issue[0] for issue in issues), "missing_field": ", ".join(issue[0] for issue in issues), "recommended_correction": " ".join(issue[1] for issue in issues)})
        else:
            if candidate["image_paths"] and not candidate["image_description"]:
                warnings.append({"question_number": candidate["number"], "question": candidate["text"], "reason": "An image was found without an image description.", "missing_field": "Image description", "recommended_correction": "Add descriptive text immediately after the image."})
            if candidate["image_paths"]:
                warnings.append({"question_number": candidate["number"], "question": candidate["text"], "reason": "Image files are preserved in the standardized DOCX, but the existing JSON schema requires hosted image URLs.", "missing_field": "Question_image URL", "recommended_correction": "Upload the preserved image files to the image service before API upload."})
            questions.append(_question_json(candidate, subject_id, created_by))
    if not candidates:
        errors.append({"question_number": "—", "question": "No questions identified", "reason": "No supported question numbering pattern was found.", "missing_field": "Question structure", "recommended_correction": "Start each question with Question 1:, Q1:, or 1. and include labelled options and an Answer: line."})
    return questions, errors, candidates, warnings


def create_standard_docx(candidates, output_path, template_path):
    """Apply the master sample's page settings to normalized question content."""
    document = Document(template_path)
    document._body.clear_content()
    for position, candidate in enumerate(candidates, start=1):
        question = document.add_paragraph()
        question.add_run(f"Question {position}: ").bold = True
        question.add_run(candidate["text"])
        for image_path in candidate.get("image_paths", []):
            if image_path and Path(image_path).is_file():
                document.add_paragraph().add_run().add_picture(image_path, width=Inches(5.5))
        for description in candidate.get("image_description", []):
            paragraph = document.add_paragraph(); paragraph.add_run(f"Image description: ").bold = True; paragraph.add_run(description)
        for label, _, option in candidate["options"]:
            document.add_paragraph(f"{label}. {option}")
        answer = next((text for label, _, text in candidate["options"] if label == candidate["answer_key"]), "")
        answer_paragraph = document.add_paragraph(); answer_paragraph.add_run("Answer: ").bold = True; answer_paragraph.add_run(f"{candidate['answer_key']}. {answer}")
        if candidate["explanation"]:
            explanation = document.add_paragraph(); explanation.add_run("Explanation:").bold = True
            for text in candidate["explanation"]:
                document.add_paragraph(text)
        if candidate["brief_explanation"]:
            explanation = document.add_paragraph(); explanation.add_run("Brief Explanation:").bold = True
        for text in candidate["brief_explanation"]:
            document.add_paragraph(text)
        if candidate["detailed_explanation"]:
            detailed = document.add_paragraph(); detailed.add_run(candidate["detailed_heading"]).bold = True
            for text in candidate["detailed_explanation"]:
                paragraph = document.add_paragraph(style="List Bullet")
                match = re.match(r"^(Option\s+[A-H](?:\s*\([^)]*\))?(?:\s+is\s+(?:correct|incorrect))?\s*:)\s*(.*)$", text, re.I)
                if match:
                    paragraph.add_run(match.group(1) + " ").bold = True; paragraph.add_run(match.group(2))
                else: paragraph.add_run(text)
        if position < len(candidates): document.add_paragraph()
    document.save(output_path)
