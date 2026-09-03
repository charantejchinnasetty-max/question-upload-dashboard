import json
import uuid
import re
from docx import Document
from datetime import datetime

# =============================
# CONFIG
# =============================
INPUT_FILE = "Test 6.docx"
OUTPUT_FILE = "Test 6output.json"
ERROR_FILE = "error_questions.json"

doc = Document(INPUT_FILE)

questions = []
error_questions = []
order = 1


# =============================
# Escape HTML safely
# =============================
def escape_html(text):
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


# =============================
# Split paragraph into lines while preserving bold
# =============================
def split_paragraph_into_lines(paragraph):
    lines = []
    current_line = ""

    for run in paragraph.runs:
        text = run.text.replace("\xa0", " ")

        if not text:
            continue

        parts = text.split("\n")

        for idx, part in enumerate(parts):
            part = escape_html(part)

            if run.bold and part.strip():
                part_html = f"<strong>{part}</strong>"
            else:
                part_html = part

            current_line += part_html

            if idx < len(parts) - 1:
                if current_line.strip():
                    lines.append(current_line.strip())
                current_line = ""

    if current_line.strip():
        lines.append(current_line.strip())

    return lines


# =============================
# Detect bullet lines
# =============================
def is_bullet_line(raw_line, paragraph_style_name=""):
    style_name = paragraph_style_name.lower() if paragraph_style_name else ""

    if "bullet" in style_name or "list" in style_name:
        return True

    if raw_line.strip().startswith(("•", "·", "-", "▪", "◦")):
        return True

    return False


# =============================
# Read normal paragraphs
# =============================
def extract_paragraph_items(paragraphs):
    items = []

    for p in paragraphs:
        if not p.text.strip():
            continue

        raw_lines = p.text.replace("\xa0", " ").split("\n")
        html_lines = split_paragraph_into_lines(p)
        style_name = p.style.name if p.style else ""

        max_len = max(len(raw_lines), len(html_lines))

        while len(raw_lines) < max_len:
            raw_lines.append("")
        while len(html_lines) < max_len:
            html_lines.append("")

        for raw, html in zip(raw_lines, html_lines):
            raw = raw.strip()
            html = html.strip()

            if raw:
                items.append({
                    "raw_text": raw,
                    "html_text": html,
                    "is_bullet": is_bullet_line(raw, style_name)
                })

    return items


# =============================
# Read table cells too
# =============================
def extract_table_items(tables):
    items = []

    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text.strip():
                        continue

                    raw_lines = p.text.replace("\xa0", " ").split("\n")
                    html_lines = split_paragraph_into_lines(p)
                    style_name = p.style.name if p.style else ""

                    max_len = max(len(raw_lines), len(html_lines))

                    while len(raw_lines) < max_len:
                        raw_lines.append("")
                    while len(html_lines) < max_len:
                        html_lines.append("")

                    for raw, html in zip(raw_lines, html_lines):
                        raw = raw.strip()
                        html = html.strip()

                        if raw:
                            items.append({
                                "raw_text": raw,
                                "html_text": html,
                                "is_bullet": is_bullet_line(raw, style_name)
                            })

    return items


# =============================
# Merge all document content
# =============================
doc_items = []
doc_items.extend(extract_paragraph_items(doc.paragraphs))
doc_items.extend(extract_table_items(doc.tables))

# Optional sort not needed — Word order is usually okay


# =============================
# HTML helpers
# =============================
def format_html(text):
    return f"<p>{text}</p>"


def format_colon_line(html):
    """
    Convert:
    Indirect Action: blah blah
    into:
    <strong>Indirect Action:</strong> blah blah
    """

    clean_html = re.sub(r'^[•·\-\▪\◦]\s*', '', html).strip()

    if ":" in clean_html:
        parts = clean_html.split(":", 1)
        left = parts[0].strip()
        right = parts[1].strip()

        if "<strong>" in left:
            return clean_html

        return f"<strong>{left}:</strong> {right}"

    return clean_html


# =============================
# Build Explanation HTML
# =============================
def build_explanation_html(explanation_items):
    if not explanation_items:
        return ""

    html_parts = []
    in_list = False

    for item in explanation_items:
        raw = item["raw_text"].strip()
        html = item["html_text"].strip()
        bullet = item["is_bullet"]

        if not raw:
            continue

        # Remove manual bullet symbols
        html = re.sub(r'^[•·\-\▪\◦]\s*', '', html).strip()

        # Auto format "Heading: content"
        html = format_colon_line(html)

        # Auto convert heading-like lines to bullets
        if ":" in raw and not raw.lower().startswith(("question", "answer", "explanation")):
            bullet = True

        if bullet:
            if not in_list:
                html_parts.append("<ul>")
                in_list = True

            html_parts.append(f"    <li>{html}</li>")
        else:
            if in_list:
                html_parts.append("</ul>")
                in_list = False

            html_parts.append(f"<p>{html}</p>")

    if in_list:
        html_parts.append("</ul>")

    return "\n".join(html_parts)


# =============================
# Main Parser
# =============================
i = 0

while i < len(doc_items):
    line = doc_items[i]["raw_text"]

    # Flexible Question detection
    if re.match(r"Question\s*\d*\s*:", line, re.IGNORECASE):

        question_text = re.sub(r"Question\s*\d*\s*:\s*", "", line, flags=re.IGNORECASE).strip()
        raw_question_block = [line]
        i += 1

        # -----------------------------
        # Extract Options dynamically
        # Supports A/B/C/D/E/F etc.
        # -----------------------------
        options = []
        option_map = {}
        option_count = 0

        while i < len(doc_items):
            current_line = doc_items[i]["raw_text"]
            raw_question_block.append(current_line)

            # Stop at Answer
            if re.match(r'Answer\s*:', current_line, re.IGNORECASE):
                break

            # Match options like A. / B) / C. / D. / E.
            if re.match(r'^[A-Z][\.\)]', current_line):
                letter = current_line[0].upper()
                option_text = re.sub(r'^[A-Z][\.\)]\s*', '', current_line).strip()

                option_count += 1
                option_map[letter] = option_count

                options.append({
                    "name": format_html(option_text),
                    "value": option_count,
                    "imgUrl": "",
                    "upload": False
                })

                i += 1
                continue

            # If not option and not Answer, stop
            break

        # -----------------------------
        # Extract Answer
        # -----------------------------
        answer = None
        answer_letter = None

        if i < len(doc_items) and re.match(r'Answer\s*:', doc_items[i]["raw_text"], re.IGNORECASE):
            match = re.search(r'Answer\s*:\s*([A-Z])', doc_items[i]["raw_text"], re.IGNORECASE)
            if match:
                answer_letter = match.group(1).upper()
                answer = option_map.get(answer_letter)

            i += 1

        # -----------------------------
        # If no answer, log error and skip
        # -----------------------------
        if answer is None:
            print("❌ Skipping (no valid answer):", question_text)

            error_questions.append({
                "question": question_text,
                "reason": f"Answer '{answer_letter}' not found in options OR missing answer line",
                "raw_block": raw_question_block
            })

            # Skip until next question
            while i < len(doc_items) and not re.match(r"Question\s*\d*\s*:", doc_items[i]["raw_text"], re.IGNORECASE):
                i += 1

            continue

        # -----------------------------
        # Extract Explanation
        # -----------------------------
        explanation_items = []

        if i < len(doc_items) and re.match(r'Explanation\s*:', doc_items[i]["raw_text"], re.IGNORECASE):

            first_line_raw = re.sub(r'^Explanation\s*:\s*', '', doc_items[i]["raw_text"], flags=re.IGNORECASE).strip()
            first_line_html = re.sub(r'^<strong>Explanation:</strong>\s*', '', doc_items[i]["html_text"], flags=re.IGNORECASE).strip()

            if first_line_raw:
                explanation_items.append({
                    "raw_text": first_line_raw,
                    "html_text": first_line_html,
                    "is_bullet": False
                })

            i += 1

            while i < len(doc_items) and not re.match(r"Question\s*\d*\s*:", doc_items[i]["raw_text"], re.IGNORECASE):
                explanation_items.append(doc_items[i])
                i += 1

        formatted_description = build_explanation_html(explanation_items)

        # -----------------------------
        # Final JSON object
        # -----------------------------
        questions.append({
            "uuid": str(uuid.uuid4()),

            "Question": format_html(question_text),

            "type": "SINGLE",

            "options": options,

            "Answer": {
                "options": answer
            },

            "Question_image": "",

            "description": formatted_description if formatted_description else "",

            "difficulty": "EASY",

            "previous_apperance": "",
            "perals": [],
            "tags": None,
            "questionId": "",

            "flags": {
                "pro": False,
                "editable": True,
                "qBank": True,
                "active": True,
                "testSeries": True,
                "paid": True
            },

            "testSeries": None,

            "Question_order": order,

            "syllabus": [
                "61dbfdcdde0d4708d95d83b2"
            ],

            "math_library": "no",

            "description_image": "",

            "createdOn": datetime.utcnow().isoformat() + "Z",

            "createdBy": {
                "uuid": "290604b4-498f-4428-8e4c-92cb9f3240d0",
                "name": "Akhil"
            },

            "modifiedOn": None
        })

        order += 1

    else:
        i += 1


# =============================
# Save Output
# =============================
with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
    json.dump(questions, f, indent=4, ensure_ascii=False)

with open(ERROR_FILE, "w", encoding="utf-8") as f:
    json.dump(error_questions, f, indent=4, ensure_ascii=False)

print("✅ Total Questions Extracted:", len(questions))
print("⚠️ Total Error Questions:", len(error_questions))
print(f"📁 Output saved to: {OUTPUT_FILE}")
print(f"📁 Error log saved to: {ERROR_FILE}")